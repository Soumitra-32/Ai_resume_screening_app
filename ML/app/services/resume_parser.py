import re
import io
import fitz  # PyMuPDF
import docx
from typing import Optional, Tuple, List

EMAIL_REGEX = re.compile(
    r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
)

PHONE_REGEX = re.compile(
    r"(?:\+?\d{1,3}[\s.-]?)?"          # optional country code, e.g. +880, +1
    r"(?:\(?\d{2,4}\)?[\s.-]?){2,4}"   # groups of 2-4 digits, 2-4 times
    r"\d{2,4}"                         # final group
)

EXPERIENCE_REGEXES = [
    re.compile(
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)"
        r"(?:\s+of)?\s+(?:professional\s+)?experience",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:experience|work experience)\s*[:\-]?\s*"
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)",
        re.IGNORECASE,
    ),
]

# Common resume section headers — used to know when the "header block" ends
SECTION_HEADERS = {
    "summary", "objective", "experience", "work experience", "education",
    "skills", "projects", "certifications", "achievements", "profile",
    "contact", "employment history", "technical skills", "references",
}

# Section headers whose content should count toward semantic-matching text
RELEVANT_SECTION_HEADERS = {
    "experience", "work experience", "employment history", "projects",
    "technical skills", "skills", "summary", "profile",
}

# Section headers whose content should be excluded from semantic-matching text
IRRELEVANT_SECTION_HEADERS = {
    "education", "certifications", "references", "achievements",
    "hobbies", "interests",
}

# Words that show up in header lines but are never part of a person's name
NON_NAME_TOKENS = {
    "resume", "cv", "curriculum", "vitae", "portfolio", "linkedin",
    "github", "phone", "email", "address", "profile",
}

# Words that indicate a line is a job title rather than a person's name
COMMON_JOB_TITLE_WORDS = {
    "engineer", "developer", "manager", "director", "analyst", "designer",
    "consultant", "specialist", "architect", "administrator", "lead",
    "intern", "coordinator", "officer", "executive", "scientist",
    "software", "senior", "junior", "full", "stack", "product", "project",
}

TITLE_CASE_WORD = re.compile(r"^[A-Z][a-zA-Z'.-]*$")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw text from a PDF using PyMuPDF."""
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from a DOCX using python-docx."""
    file_stream = io.BytesIO(file_bytes)
    document = docx.Document(file_stream)

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n".join(parts).strip()


def extract_relevant_text_for_matching(text: str) -> str:
    """
    Returns a subset of the resume focused on experience/skills content,
    for use in semantic similarity scoring — excludes sections like
    education, certifications, and hobbies that shouldn't drive the
    embedding-based match score.
    Falls back to the full text if no section headers are detected.
    """
    lines = text.split("\n")
    relevant_lines = []
    current_section_relevant = True  # include header-less leading content by default

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower()

        if lower in RELEVANT_SECTION_HEADERS:
            current_section_relevant = True
            continue
        if lower in IRRELEVANT_SECTION_HEADERS:
            current_section_relevant = False
            continue
        if lower in SECTION_HEADERS:
            # A section header we don't have a strong opinion on — keep it, be permissive
            current_section_relevant = True
            continue

        if current_section_relevant and stripped:
            relevant_lines.append(stripped)

    result = "\n".join(relevant_lines).strip()
    return result if result else text  # fallback: better to use everything than nothing


def extract_email(text: str) -> Optional[str]:
    match = EMAIL_REGEX.search(text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    match = PHONE_REGEX.search(text)
    if not match:
        return None
    phone = match.group(0).strip()
    digits_only = re.sub(r"\D", "", phone)
    # BD mobile numbers are 11 digits locally (01XXXXXXXXX) or 13 with
    # country code (+8801XXXXXXXXX); loosen the bound accordingly.
    if len(digits_only) < 7 or len(digits_only) > 15:
        return None
    return phone


def extract_experience_years(text: str) -> Optional[float]:
    """
    Extract years of experience, preferring an explicit summary/profile
    statement near the top of the resume over scanning the whole document
    (which can pick up project durations or unrelated year references).
    """
    lines = text.split("\n")
    # Look only in the first ~40 lines (roughly the header/summary section)
    head_text = "\n".join(lines[:40])

    matches = []
    for pattern in EXPERIENCE_REGEXES:
        for match in pattern.finditer(head_text):
            try:
                matches.append(float(match.group(1)))
            except ValueError:
                pass

    if matches:
        return max(matches)

    # Fall back to scanning the full document only if nothing was found up top
    for pattern in EXPERIENCE_REGEXES:
        for match in pattern.finditer(text):
            try:
                matches.append(float(match.group(1)))
            except ValueError:
                pass

    return max(matches) if matches else None


def _looks_like_name(line: str) -> bool:
    """Heuristic check: is this line plausibly a person's full name?"""
    line = line.strip()
    if not line or len(line) > 60:
        return False
    if "@" in line or any(c.isdigit() for c in line):
        return False
    if any(ch in line for ch in ("|", "•", "http", "www", "/")):
        return False

    lower = line.lower()
    if lower in SECTION_HEADERS:
        return False
    words_lower = lower.split()
    if any(tok in words_lower for tok in NON_NAME_TOKENS):
        return False
    # Reject lines that look like a job title rather than a person's name
    if any(tok in words_lower for tok in COMMON_JOB_TITLE_WORDS):
        return False

    words = line.split()
    if not (1 <= len(words) <= 4):
        return False

    # Most words in a name line should be Title-Case (allow ALL CAPS names too)
    title_like = sum(
        1 for w in words
        if TITLE_CASE_WORD.match(w) or w.isupper()
    )
    return title_like >= max(1, len(words) - 1)


def extract_name(text: str) -> Optional[str]:
    """
    Heuristic name extraction (no NLP model required):
    Scan the first ~10 non-empty lines and return the first one that
    looks like a plausible person name and isn't a section header/contact line.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:10]:
        if _looks_like_name(line):
            # Normalize ALL CAPS names to Title Case for readability
            if line.isupper():
                return line.title()
            return line
    return None


def extract_entities(text: str, max_chars: int = 5000) -> List[str]:
    """
    Lightweight heuristic entity spotting (no NLP model):
    - Lines matching common section headers
    - Capitalized multi-word phrases that look like organizations
    This is intentionally simple; it's a placeholder until Phase 4
    where embeddings-based scoring adds real semantic understanding.
    """
    entities = []
    seen = set()
    snippet = text[:max_chars]

    for line in snippet.split("\n"):
        line = line.strip()
        lower = line.lower()
        if lower in SECTION_HEADERS and line not in seen:
            seen.add(line)
            entities.append(f"{line} (SECTION)")

    # Naive "Organization-like" phrase detector: 2-4 Title-Case words in a row,
    # not already flagged as a section header or the resume owner's name line.
    org_pattern = re.compile(r"\b([A-Z][a-zA-Z&.,-]*(?:\s+[A-Z][a-zA-Z&.,-]*){1,3})\b")
    for match in org_pattern.finditer(snippet):
        phrase = match.group(1).strip()
        if phrase not in seen and len(phrase.split()) <= 4:
            seen.add(phrase)
            entities.append(f"{phrase} (ORG_CANDIDATE)")
        if len(entities) >= 30:
            break

    return entities


def parse_resume_file(
    file_bytes: bytes, file_extension: str
) -> Tuple[str, Optional[str], Optional[str], Optional[str], Optional[float], List[str]]:
    """
    Main orchestration function.
    Returns: (raw_text, name, email, phone, experience_years, entities)
    """
    if file_extension == ".pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif file_extension == ".docx":
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")

    if not raw_text:
        raw_text = ""

    name = extract_name(raw_text)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    experience_years = extract_experience_years(raw_text)
    entities = extract_entities(raw_text)

    return raw_text, name, email, phone, experience_years, entities


def get_resume_text(
    file_bytes: Optional[bytes] = None,
    file_extension: Optional[str] = None,
    raw_text: Optional[str] = None,
) -> str:
    """
    Unified entry point that works with EITHER:
    - a raw file (PDF/DOCX bytes) -> extracts text
    - already-extracted text (e.g. from a CSV column like Resume_str)

    Exactly one of (file_bytes + file_extension) OR raw_text should be provided.
    Never throws on missing/empty input — always returns a string (possibly empty).
    """
    if raw_text is not None:
        return str(raw_text).strip()

    if file_bytes is not None and file_extension is not None:
        ext = file_extension.lower()
        try:
            if ext == ".pdf":
                return extract_text_from_pdf(file_bytes)
            elif ext == ".docx":
                return extract_text_from_docx(file_bytes)
            else:
                return ""
        except Exception:
            return ""

    return ""